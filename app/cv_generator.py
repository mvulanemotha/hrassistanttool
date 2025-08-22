import io
import os
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI
from dotenv import load_dotenv
import tempfile
from docx2pdf import convert
from fastapi import UploadFile, File, Form, HTTPException, Depends
from sqlalchemy.orm import Session
import base64
from fastapi.responses import JSONResponse
from docx.table import Table
from docx.text.paragraph import Paragraph
from typing import Dict, List, Tuple
import html


# --- Load environment variables ---
load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("❌ Missing GROQ_API_KEY in .env file")

client = OpenAI(
    api_key=GROQ_API_KEY,
    base_url="https://api.groq.com/openai/v1"
)

# --- Enhanced Helper functions ---

def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract clean text from DOCX, including paragraphs and tables."""
    doc = Document(io.BytesIO(docx_bytes))
    parts = []

    # Normal paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    # Text inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    return "\n\n".join(parts)

def generate_attorney_cv_with_llm(user_cv_text: str, template_text: str) -> str:
    """Generate professionally formatted attorney CV text"""
    prompt = f"""
    You are an expert legal resume writer. Rewrite the user's attorney CV using legal industry best practices.
    
    Important guidelines for an attorney resume:
    1. Focus on legal expertise, case work, and specialized knowledge
    2. Use strong legal action verbs (litigated, negotiated, drafted, advised, represented)
    3. Highlight specific legal achievements and case outcomes
    4. Include bar admissions, court admissions, and certifications
    5. Emphasize relevant legal experience and specialized practice areas
    6. Use professional legal terminology appropriate for the field
    7. Structure content for a two-column legal resume format
    
    --- USER CV CONTENT ---
    {user_cv_text}
    
    --- TEMPLATE STYLE ---
    {template_text}
    
    Now produce an enhanced, professional attorney CV:
    """
    
    try:
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=3000
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error generating attorney CV with LLM: {str(e)}")
        return create_fallback_attorney_cv(user_cv_text)

def create_fallback_attorney_cv(user_cv_text: str) -> str:
    """Create a basic formatted attorney CV if LLM generation fails"""
    # Parse the user's CV text into sections
    sections = {
        'contact': [],
        'profile': [],
        'experience': [],
        'education': [],
        'skills': [],
        'interests': []
    }
    
    current_section = None
    lines = user_cv_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Detect section headers
        line_upper = line.upper()
        if 'CONTACT' in line_upper or 'ADDRESS' in line_upper or 'PHONE' in line_upper:
            current_section = 'contact'
        elif 'PROFILE' in line_upper or 'SUMMARY' in line_upper or 'OBJECTIVE' in line_upper:
            current_section = 'profile'
        elif 'EXPERIENCE' in line_upper or 'EMPLOYMENT' in line_upper:
            current_section = 'experience'
        elif 'EDUCATION' in line_upper or 'DEGREE' in line_upper or 'JURIS' in line_upper:
            current_section = 'education'
        elif 'SKILL' in line_upper or 'EXPERTISE' in line_upper or 'ADMISSION' in line_upper:
            current_section = 'skills'
        elif 'INTEREST' in line_upper or 'ACTIVIT' in line_upper:
            current_section = 'interests'
        elif current_section:
            sections[current_section].append(line)
    
    # Build the formatted CV
    formatted_cv = "CONTACT INFORMATION\n"
    formatted_cv += "\n".join(sections['contact']) + "\n\n"
    
    formatted_cv += "PROFILE\n"
    formatted_cv += "\n".join(sections['profile']) + "\n\n"
    
    formatted_cv += "EXPERIENCE\n"
    formatted_cv += "\n".join(sections['experience']) + "\n\n"
    
    formatted_cv += "EDUCATION\n"
    formatted_cv += "\n".join(sections['education']) + "\n\n"
    
    formatted_cv += "KEY SKILLS\n"
    formatted_cv += "\n".join(sections['skills']) + "\n\n"
    
    formatted_cv += "INTERESTS\n"
    formatted_cv += "\n".join(sections['interests'])
    
    return formatted_cv

def create_attorney_docx_from_text(generated_text: str) -> bytes:
    """Create a two-column attorney resume DOCX based on the generated text"""
    doc = Document()
    
    # Set up document formatting
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Parse the generated text into sections
    sections = parse_attorney_sections(generated_text)
    
    # Create the main table with 2 columns
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Set column widths (approximately 40%/60% split)
    for cell in table.columns[0].cells:
        cell.width = Inches(2.5)
    for cell in table.columns[1].cells:
        cell.width = Inches(4.0)
    
    # Get the cells
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    
    # Add content to left cell (Contact, Education, Skills, Interests)
    add_left_column_content(left_cell, sections)
    
    # Add content to right cell (Profile, Experience)
    add_right_column_content(right_cell, sections)
    
    # Save to bytes
    output_stream = io.BytesIO()
    doc.save(output_stream)
    return output_stream.getvalue()

def parse_attorney_sections(generated_text: str) -> Dict[str, List[str]]:
    """Parse the generated text into sections for an attorney resume"""
    sections = {
        'contact': [],
        'profile': [],
        'experience': [],
        'education': [],
        'skills': [],
        'interests': []
    }
    
    current_section = None
    lines = generated_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Detect section headers
        line_upper = line.upper()
        if 'CONTACT' in line_upper:
            current_section = 'contact'
            continue
        elif 'PROFILE' in line_upper or 'SUMMARY' in line_upper:
            current_section = 'profile'
            continue
        elif 'EXPERIENCE' in line_upper or 'EMPLOYMENT' in line_upper:
            current_section = 'experience'
            continue
        elif 'EDUCATION' in line_upper:
            current_section = 'education'
            continue
        elif 'SKILL' in line_upper or 'EXPERTISE' in line_upper:
            current_section = 'skills'
            continue
        elif 'INTEREST' in line_upper or 'ACTIVIT' in line_upper:
            current_section = 'interests'
            continue
            
        # Add content to the current section
        if current_section and line:
            sections[current_section].append(line)
    
    return sections

def add_left_column_content(cell, sections: Dict[str, List[str]]):
    """Add content to the left column of the attorney resume"""
    # Add contact information
    if sections['contact']:
        for line in sections['contact']:
            p = cell.add_paragraph(line)
            p.paragraph_format.space_after = Pt(3)
        cell.add_paragraph()
    
    # Add education
    if sections['education']:
        p = cell.add_paragraph()
        run = p.add_run("EDUCATION")
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        
        for line in sections['education']:
            p = cell.add_paragraph(line)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Pt(0)
        cell.add_paragraph()
    
    # Add skills
    if sections['skills']:
        p = cell.add_paragraph()
        run = p.add_run("KEY SKILLS")
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        
        for line in sections['skills']:
            p = cell.add_paragraph(line)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Pt(0)
        cell.add_paragraph()
    
    # Add interests
    if sections['interests']:
        p = cell.add_paragraph()
        run = p.add_run("INTERESTS")
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        
        for line in sections['interests']:
            p = cell.add_paragraph(line)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Pt(0)

def add_right_column_content(cell, sections: Dict[str, List[str]]):
    """Add content to the right column of the attorney resume"""
    # Add profile
    if sections['profile']:
        p = cell.add_paragraph()
        run = p.add_run("PROFILE")
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        
        for line in sections['profile']:
            p = cell.add_paragraph(line)
            p.paragraph_format.space_after = Pt(3)
        cell.add_paragraph()
    
    # Add experience
    if sections['experience']:
        p = cell.add_paragraph()
        run = p.add_run("EXPERIENCE")
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        
        for line in sections['experience']:
            p = cell.add_paragraph(line)
            p.paragraph_format.space_after = Pt(3)

def convert_docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """Convert DOCX to PDF (requires MS Word)"""
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "temp.docx")
        pdf_path = os.path.join(tmpdir, "temp.pdf")

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        convert(docx_path, pdf_path)

        with open(pdf_path, "rb") as f:
            return f.read()